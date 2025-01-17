from docx import Document
from datetime import datetime
import os
import pandas as pd
import logging
from utils import validate_file, format_dates, clean_file_name, load_word_template, logging_error,replace_placeholders_in_cell,replace_placeholders_in_paragraph






def generate_certificates(excel_file_path, word_template_path, output_folder,
                          broj_certifkata_ui, naziv_firme_ui, adresa_firme_ui, datum_dokumenta_ui, grad_ui):
    try:
        # Load Excel data
        data = pd.read_excel(excel_file_path)

        # Format date columns
        for column in data.columns:
            if pd.api.types.is_datetime64_any_dtype(data[column]):
                data[column] = data[column].dt.strftime('%d.%m.%Y')

        current_year = datetime.now().year

        # Iterate over rows in the Excel file
        for index, row in data.iterrows():
            try:
                # Load and prepare the Word template
                doc = Document(word_template_path)
                ime_prezime = f"{row.get('ime', '').strip()}_{row.get('prezime', '').strip()}"
                safe_file_name = clean_file_name(ime_prezime)

                # Prepare placeholders
                placeholders = {
                    "<<broj_certifkata>>": f"{broj_certifkata_ui}-{index + 1}/{current_year}",
                    "<<datum_dokumenta>>": datum_dokumenta_ui,
                    "<<naziv_firme>>": naziv_firme_ui,
                    "<<adresa_firme>>": adresa_firme_ui,
                    "<<grad_firme>>": grad_ui,
                    "<<ime_prezime>>": f"{row.get('ime', '').strip()} {row.get('prezime', '').strip()}",
                    "<<datum_rodjenja>>": row.get('datum_rodjenja', 'N/A') if isinstance(row.get('datum_rodjenja'), str)
                                        else row.get('datum_rodjenja').strftime('%d.%m.%Y') if pd.notna(row.get('datum_rodjenja')) else 'N/A',
                    "<<adresa_stanovanja>>": row.get('adresa_stanovanja', 'N/A'),
                    "<<grad_stanovanja>>": row.get('grad_stanovanja', 'N/A'),
                    "<<vrsta_posla>>": row.get('vrsta_posla', 'N/A').strip()
                }

                # Replace placeholders in paragraphs
                for paragraph in doc.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, placeholders)

                # Replace placeholders in table cells
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_placeholders_in_cell(cell, placeholders)

                # Save the updated document
                output_file = os.path.join(output_folder, f"Certifikat_{safe_file_name}.docx")
                doc.save(output_file)
                logging.info(f"Certificate generated: {output_file}")

            except Exception as e:
                logging.error(f"Error processing row {index + 1}: {e}")

    except Exception as e:
        logging.error(f"Failed to generate certificates: {e}")

