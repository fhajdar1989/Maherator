from docx import Document
from datetime import datetime
import pandas as pd
import os
import logging
from utils import validate_file, replace_placeholders_in_cell, replace_placeholders_in_paragraph

def format_dates(df):
    """
    Converts the 'datum_rodjenja' column to a consistent date format (dd.mm.yyyy).
    Handles mixed types and invalid values gracefully.
    """
    try:
        if 'datum_rodjenja' in df.columns:
            def parse_date(value):
                try:
                    # Attempt parsing as a date
                    if isinstance(value, str):
                        return pd.to_datetime(value, errors='coerce', dayfirst=True).strftime('%d.%m.%Y')
                    elif isinstance(value, (datetime, pd.Timestamp)):
                        return value.strftime('%d.%m.%Y')
                    else:
                        return "N/A"
                except Exception:
                    return "N/A"

            df['datum_rodjenja'] = df['datum_rodjenja'].apply(parse_date)
        return df
    except Exception as e:
        logging.error(f"Error formatting dates: {e}")
        raise

def populate_table(table, data, static_placeholders, table_type="first"):
    """
    Populates placeholders in a table dynamically.
    """
    try:
        if table_type == "first":
            template_row = table.rows[1]
            for index, row_data in data.iterrows():
                placeholders = {
                    "<<redni_broj>>": str(index + 1),
                    "<<ime>>": row_data.get("ime", "").strip(),
                    "<<ocevo_ime>>": row_data.get("ocevo_ime", "").strip(),
                    "<<prezime>>": row_data.get("prezime", "").strip(),
                    "<<datum_rodjenja>>": row_data.get("datum_rodjenja", "N/A"),
                    "<<naziv_firme>>": static_placeholders["<<naziv_firme>>"],
                    "<<adresa_firme>>": static_placeholders["<<adresa_firme>>"],
                    "<<grad_firme>>": static_placeholders["<<grad_firme>>"],
                }
                new_row = table.add_row()
                for i, cell in enumerate(new_row.cells):
                    cell.text = template_row.cells[i].text
                for cell in new_row.cells:
                    replace_placeholders_in_cell(cell, placeholders)
            table._tbl.remove(template_row._element)

        elif table_type == "second":
            template_row_1, template_row_2 = table.rows[:2]
            common_placeholder = data.iloc[0].get("vrsta_posla", "N/A").strip()
            first_row = table.add_row().cells[0]
            first_row.text = f"BEZBJEDAN I SIGURAN RAD NA RADNOM MJESTU {common_placeholder}"

            for index, row_data in data.iterrows():
                placeholders = {
                    "<<broj_certifkata>>": static_placeholders["<<broj_certifkata>>"],
                    "<<redni_broj>>": str(index + 1),
                    "<<godina>>": static_placeholders["<<godina>>"],
                }
                second_row = table.add_row().cells[0]
                second_row.text = f"te se izdaje UVJERENJE br. {placeholders['<<broj_certifkata>>']}-{placeholders['<<redni_broj>>']}/{placeholders['<<godina>>']}"
            table._tbl.remove(template_row_1._element)
            table._tbl.remove(template_row_2._element)

    except Exception as e:
        logging.error(f"Error populating table ({table_type}): {e}")
        raise

def generate_agreements(excel_file_path, word_template_path, output_folder,
                        broj_certifkata, naziv_firme, adresa_firme, grad_firme, datum_dokumenta):
    try:
        validate_file(excel_file_path, "Excel file")
        validate_file(word_template_path, "Word template")

        if not os.path.isdir(output_folder):
            raise FileNotFoundError(f"The specified output folder does not exist: {output_folder}")

        data = format_dates(pd.read_excel(excel_file_path))
        current_year = datetime.now().year

        doc = Document(word_template_path)  # Directly load template without conversion
        broj_dokumenta = f"{broj_certifkata}/{current_year}"

        static_placeholders = {
            "<<broj_dokumenta>>": broj_dokumenta,
            "<<broj_certifkata>>": broj_certifkata,
            "<<naziv_firme>>": naziv_firme,
            "<<adresa_firme>>": adresa_firme,
            "<<grad_firme>>": grad_firme,
            "<<datum_dokumenta>>": datum_dokumenta,
            "<<godina>>": str(current_year),
        }

        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, static_placeholders)

        first_table = next((tbl for tbl in doc.tables if "<<redni_broj>>" in tbl.rows[1].cells[0].text), None)
        if first_table:
            logging.info("First table found. Populating...")
            populate_table(first_table, data, static_placeholders, table_type="first")
        else:
            logging.warning("First table not found.")

        second_table = next((tbl for tbl in doc.tables if any("<<vrsta_posla>>" in cell.text for row in tbl.rows for cell in row.cells)), None)
        if second_table:
            logging.info("Second table found. Populating...")
            populate_table(second_table, data, static_placeholders, table_type="second")
        else:
            logging.warning("Second table not found.")

        formatted_datum_dokumenta = datum_dokumenta.replace('.', '_')
        output_file_path = os.path.join(output_folder, f"Zapisnik_{formatted_datum_dokumenta}.docx")

        doc.save(output_file_path)
        if os.path.exists(output_file_path):
            logging.info(f"Agreement document successfully saved: {output_file_path}")
        else:
            logging.error(f"Failed to save the agreement document: {output_file_path}")

    except Exception as e:
        logging.error(f"An unexpected error occurred: {e}")
        raise
