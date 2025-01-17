import logging
import os
import pandas as pd
from docx import Document
import ctypes
import win32com.client as win32
import webview
# Configure logging
logging.basicConfig(
    filename="application.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)

def validate_file(file_path, file_type):
    """Check if a file exists and is accessible."""
    if not os.path.exists(file_path):
        logging.error(f"{file_type} file not found: {file_path}")
        raise FileNotFoundError(f"{file_type} file not found: {file_path}")
    logging.info(f"{file_type} file validated: {file_path}")


def format_dates(df):
    """
    Converts 'datum_rodjenja' to a consistent date format and ensures all date columns are valid.
    """
    try:
        if 'datum_rodjenja' in df.columns:
            df['datum_rodjenja'] = pd.to_datetime(
                df['datum_rodjenja'], errors='coerce', dayfirst=True
            ).dt.strftime('%d.%m.%Y')
        return df
    except Exception as e:
        logging.error(f"Error formatting dates: {e}")
        raise


def clean_file_name(name):
    """
    Ensures the file name is safe for the filesystem by removing or replacing unsafe characters.
    """
    import re
    return re.sub(r'[\\/*?:"<>|]', "_", name)




def convert_doc_to_docx(doc_file_path):
    """
    Converts a .doc file to .docx format using Microsoft Word.
    """
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(doc_file_path)
        docx_file_path = doc_file_path + "x"  # Add 'x' to save as .docx
        doc.SaveAs2(docx_file_path, FileFormat=16)  # FileFormat 16 is for .docx
        doc.Close()
        logging_error(f"Converted {doc_file_path} to {docx_file_path}")
        return docx_file_path
    except Exception as e:
        logging_error(f"Error converting {doc_file_path} to .docx: {e}")
        return None
    finally:
        word.Quit()
        del word  # Ensure the COM object is released


def load_word_template(file_path):
    """
    Loads a Word template, converting .doc to .docx if necessary.
    """
    if file_path.endswith(".doc"):
        logging_error(f"Converting .doc file: {file_path}")
        file_path = convert_doc_to_docx(file_path)
        if not file_path:
            raise ValueError("Conversion from .doc to .docx failed.")
    return Document(file_path)



def logging_error(message, is_error=True):
    """
    Logs an error message to both the console and a log file if it's an error.
    Normal application behaviors (like exits) won't be logged.
    
    Args:
        message (str): The message to log.
        is_error (bool): Whether the message is an error. Defaults to True.
    """
    if is_error:
        print(message)
        with open("error.log", "a") as log_file:
            log_file.write(f"{message}\n")


def log_and_alert_error(e):
    error_message = f"Error: {str(e)}"
    logging_error(error_message)
    try:
        webview.windows[0].evaluate_js(f'showResult("{error_message}")')
    except webview.errors.JavascriptException:
        print(f"Failed to call showResult in JavaScript. Error: {error_message}")



def set_app_icon(icon_path):
    """
    Set the application icon for the taskbar (Windows-specific).
    """
    if os.name == "nt":
        myappid = 'Maherator.App'
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)


def replace_placeholders_in_cell(cell, placeholders):
    """
    Replaces placeholders in all paragraphs within a table cell.
    """
    for paragraph in cell.paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        for placeholder, replacement in placeholders.items():
            full_text = full_text.replace(placeholder, replacement)
        for run in paragraph.runs:
            run.text = ""  # Clear the existing text
        paragraph.add_run(full_text)  # Add the updated text


def replace_placeholders_in_paragraph(paragraph, placeholders):
    """
    Replaces placeholders in a paragraph, preserving formatting and handling split runs.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    for placeholder, replacement in placeholders.items():
        full_text = full_text.replace(placeholder, replacement)
    for run in paragraph.runs:
        run.text = ""
    paragraph.add_run(full_text)
